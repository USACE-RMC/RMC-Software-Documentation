from docx import Document
from docx.text.paragraph import Paragraph
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from PIL import Image
from io import BytesIO
import os
import re
import hashlib


def iter_block_items(doc):
    # Yields paragraphs and tables in order
    for block in doc.element.body:
        if block.tag.endswith("tbl"):
            yield block
        elif block.tag.endswith("p"):
            yield Paragraph(block, doc)


def extract_figure_images(docx_path, fig_img_output_path, extract_figures):
    """
    Walks captions (style contains 'caption' and text startswith 'Figure'),
    backtracks to the preceding image, and:
      - saves embedded images as figures/figureN.png
      - reserves figureN.png for linked/missing images (writes 1×1 placeholder if extract_figures)
    Returns maps for inline usage plus a list-aligned figure map.
    """
    os.makedirs(fig_img_output_path, exist_ok=True)
    doc = Document(docx_path)

    figures_dir = os.path.join(fig_img_output_path, "figures")
    inline_dir = os.path.join(fig_img_output_path, "inline-images")
    os.makedirs(figures_dir, exist_ok=True)
    os.makedirs(inline_dir, exist_ok=True)

    image_hash_map = {}
    inline_image_hash_map = {}
    inline_image_locations = []  # (para_idx, run_idx, filename)
    figure_image_hashes = set()

    # Build relationship maps (embedded vs external)
    rels_embedded = {}  # rId -> bytes
    rels_external = {}  # rId -> external target
    for rel in doc.part.rels.values():
        if rel.reltype == RT.IMAGE:
            if getattr(rel, "is_external", False):
                rels_external[rel.rId] = rel.target_ref
            else:
                rels_embedded[rel.rId] = rel.target_part.blob

    def _get_rEmbed_from_run(run):
        # python-docx's xpath already knows the "a" prefix
        blips = run._element.xpath(".//a:blip")
        if blips:
            return blips[0].get(qn("r:embed"))  # safer than hardcoding the full URI
        return None

    # Pass 1: captioned figures — ALWAYS reserve the slot
    figures_list = []  # aligned 1:1 with captions
    figure_count = 0

    for i, para in enumerate(doc.paragraphs):
        style = (para.style.name or "") if para.style else ""
        text = para.text.strip().replace("\xa0", " ")
        if "caption" in style.lower() and text.startswith("Figure"):
            fig_number = figure_count + 1
            image_data = None
            external_target = None

            # search backward for the image preceding this caption
            for prev_para_idx in range(i - 1, -1, -1):
                prev_para = doc.paragraphs[prev_para_idx]
                for run_idx in reversed(range(len(prev_para.runs))):
                    run = prev_para.runs[run_idx]
                    if "graphic" not in run._element.xml:
                        continue
                    rId = _get_rEmbed_from_run(run)
                    if not rId:
                        continue
                    if rId in rels_embedded:
                        image_data = rels_embedded[rId]
                        break
                    elif rId in rels_external:
                        external_target = rels_external[rId]
                        # keep scanning in case there’s also an embedded image
                if image_data or external_target:
                    break

            filename = f"figure{fig_number}.png"
            save_path = os.path.join(figures_dir, filename)

            status = "embedded"
            if image_data is not None:
                # save embedded image
                try:
                    if extract_figures:
                        img = Image.open(BytesIO(image_data))
                        img.save(save_path)
                    image_hash = hashlib.sha256(image_data).hexdigest()
                    image_hash_map.setdefault(image_hash, filename)
                    figure_image_hashes.add(image_hash)
                except Exception as e:
                    print(f"Warning: could not save figure {fig_number}: {e}")
                    status = "error"
            else:
                # linked or missing — reserve slot, optionally write a placeholder
                status = "external" if external_target else "missing"
                if extract_figures:
                    try:
                        ph = Image.new("RGBA", (1, 1), (0, 0, 0, 0))
                        ph.save(save_path)
                    except Exception as e:
                        print(f"Warning: could not create placeholder for figure {fig_number}: {e}")
                        status = "error"

            # Build caption text without the "Figure N:" prefix (if you need it later)
            caption = re.sub(r"^Figure\s*\d+[\.:]*\s*", "", text).lstrip()
            figures_list.append(
                {
                    "figure_number": fig_number,
                    "figKey": f"figure-{fig_number}",
                    "src": f"{os.path.basename(fig_img_output_path)}\\figures\\{filename}",
                    "alt": caption,
                    "caption": caption,
                    "status": status,
                    **({"external_target": external_target} if external_target else {}),
                }
            )

            figure_count += 1

    # Pass 2: inline images (only embedded; skip ones already used for figures)
    for para_idx, para in enumerate(doc.paragraphs):
        for run_idx, run in enumerate(para.runs):
            if "graphic" not in run._element.xml:
                continue
            rId = _get_rEmbed_from_run(run)
            if not rId:
                continue
            if rId in rels_embedded:
                data = rels_embedded[rId]
                img_hash = hashlib.sha256(data).hexdigest()
                if img_hash in figure_image_hashes:
                    continue
                fname = f"inline-{img_hash[:8]}.png"
                save_path = os.path.join(inline_dir, fname)
                if img_hash not in inline_image_hash_map:
                    if extract_figures:
                        try:
                            Image.open(BytesIO(data)).save(save_path)
                        except Exception as e:
                            print(
                                f"Warning: could not save inline image at ({para_idx},{run_idx}): {e}"
                            )
                            continue
                    inline_image_hash_map[img_hash] = fname
                inline_image_locations.append((para_idx, run_idx, fname))

    print(f"\n{figure_count} figure image(s) extracted/reserved.")
    print(f"{len(inline_image_locations)} inline image(s) extracted (or mapped).")

    return {
        "figure_hash_map": image_hash_map,
        "inline_image_hash_map": inline_image_hash_map,
        "inline_image_locations": inline_image_locations,
        # Add this so process_docx can consume aligned figures if you prefer
        "figures_list": figures_list,
    }


def extract_figures(doc, fig_src_path):
    figures = []
    figure_count = 0
    for i, para in enumerate(doc.paragraphs):
        style = para.style.name
        if style == "RMC_Figure":
            figure_count += 1
            figKey = f"figure-{figure_count}"
            src = f"{fig_src_path}\\figures\\figure{figure_count}.png"
            caption_paragraph = doc.paragraphs[i + 1]
            caption = caption_paragraph.text.strip()
            caption = re.sub(r"^Figure\s*\d*[\.:]*\s*", "", caption).lstrip()
            alt_text = caption
            figures.append(
                {
                    "figure_number": figure_count,
                    "figKey": figKey,
                    "src": src,
                    "alt": alt_text,
                    "caption": caption,
                }
            )
    return figures


def handle_figure_references(text):
    def replacer(match):
        figure_number = int(match.group(1))
        figKey = f"figure-{str(figure_number)}"
        if figKey:
            return f'<FigReference figKey="{figKey}" />'
        else:
            return match.group(0)

    return re.sub(r"Figure (\d+)", replacer, text)
