from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
import re


def iter_block_items(parent):
    from docx.document import Document as _Document

    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    else:
        raise ValueError("Expected a Document object")
    for child in parent_elm.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def get_true_ncols(table):
    tblGrid = table._tbl.tblGrid
    if tblGrid is not None:
        # print(f'tblGrid found with {len(tblGrid.gridCol_lst)} columns')
        return len(tblGrid.gridCol_lst)
    # Fallback if tblGrid is missing
    print("Warning: tblGrid is missing; using max number of cells in rows")
    return max(len(row.cells) for row in table.rows)


def build_table_grid(table):
    n_rows = len(table.rows)
    n_cols = get_true_ncols(table)
    # print(f"Table has {n_rows} rows and {n_cols} columns")
    grid = []
    for r_idx, row in enumerate(table.rows):
        row_data = []
        for cell in row.cells:
            text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            style = cell.paragraphs[0].style.name if cell.paragraphs else ""
            cell_info = {"text": text, "style": style, "rowSpan": 1, "colSpan": 1}
            row_data.append(cell_info)
        # Pad the row if it's short (shouldn't happen if n_cols is correct, but just in case)
        while len(row_data) < n_cols:
            row_data.append(None)
        grid.append(row_data)
    return grid


def set_spans_from_grid(grid):
    n_rows = len(grid)
    n_cols = len(grid[0]) if n_rows else 0
    for r in range(n_rows):
        for c in range(n_cols):
            cell = grid[r][c]
            if cell is None:
                continue
            # If cell above or to the left has the same value, this is a spanned cell
            if (
                r > 0
                and grid[r - 1][c] is not None
                and grid[r - 1][c]["text"] == cell["text"]
                and grid[r - 1][c]["style"] == cell["style"]
            ) or (
                c > 0
                and grid[r][c - 1] is not None
                and grid[r][c - 1]["text"] == cell["text"]
                and grid[r][c - 1]["style"] == cell["style"]
            ):
                grid[r][c] = None
                continue
            # Count colspan
            colspan = 1
            for cc in range(c + 1, n_cols):
                next_cell = grid[r][cc]
                if (
                    next_cell is not None
                    and next_cell["text"] == cell["text"]
                    and next_cell["style"] == cell["style"]
                ):
                    colspan += 1
                else:
                    break
            # Count rowspan
            rowspan = 1
            for rr in range(r + 1, n_rows):
                next_cell = grid[rr][c]
                if (
                    next_cell is not None
                    and next_cell["text"] == cell["text"]
                    and next_cell["style"] == cell["style"]
                ):
                    rowspan += 1
                else:
                    break
            cell["colSpan"] = colspan
            cell["rowSpan"] = rowspan


def print_grid(grid):
    for r, row in enumerate(grid):
        row_str = []
        for c, cell in enumerate(row):
            if cell is None:
                row_str.append("None")
            else:
                row_str.append(f"{cell['text']}[{cell.get('rowSpan',1)},{cell.get('colSpan',1)}]")
        print(" | ".join(row_str))
    print()


def extract_tables(doc):
    tables = {}
    table_index = 0
    table_number = 1
    last_caption = None
    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            style_name = getattr(getattr(block, "style", None), "name", "")
            if style_name == "RMC_Table Caption" and block.text.strip():
                last_caption = block.text.strip()
                last_caption = re.sub(r"^Table\s*\d*[\.:]*\s*", "", last_caption).lstrip()
            else:
                last_caption = None
        elif isinstance(block, Table):
            if not last_caption:
                continue
            table_data = build_table_grid(block)
            set_spans_from_grid(table_data)
            # Uncomment to debug grid visually
            # print_grid(table_data)
            if not table_data or not any(any(cell for cell in row if cell) for row in table_data):
                last_caption = None
                continue

            # Uncomment to debug grid visually
            # print_grid(table_data)

            has_spanning_cells = False
            seen_cells = set()
            for row in table_data:
                for cell in row:
                    if cell and id(cell) not in seen_cells:
                        seen_cells.add(id(cell))
                        if cell.get("rowSpan", 1) > 1 or cell.get("colSpan", 1) > 1:
                            has_spanning_cells = True

            def cell_to_obj(cell):
                if not cell:
                    return None
                # Include rowSpan and colSpan for accurate Table rendering
                return {
                    "value": cell["text"],
                    "rowSpan": cell.get("rowSpan", 1),
                    "colSpan": cell.get("colSpan", 1),
                }

            first_row = table_data[0]
            first_col = [row[0] for row in table_data if row]
            first_row_header = all(
                (cell is None or cell["style"] == "RMC_Table Header") for cell in first_row
            )
            first_col_header = all(
                (cell is None or cell["style"] == "RMC_Table Header") for cell in first_col
            )
            if first_row_header and not first_col_header:
                table_type = "TableVertical"
                header_rows = 0
                for row in table_data:
                    if all(cell is None or cell["style"] == "RMC_Table Header" for cell in row):
                        header_rows += 1
                    else:
                        break
                headers = [[cell_to_obj(cell) for cell in row] for row in table_data[:header_rows]]
                body_rows = table_data[header_rows:]
                num_cols = max(len(row) for row in body_rows)
                columns = [
                    [cell_to_obj(row[i]) if i < len(row) else None for row in body_rows]
                    for i in range(num_cols)
                ]
                data = columns
            elif first_col_header and not first_row_header:
                table_type = "TableHorizontal"

                def to_html_string(cell):
                    return cell.get("value", "")

                header_rows = 0
                for row in table_data:
                    if any(cell and cell["style"] == "RMC_Table Header" for cell in row):
                        header_rows += 1
                    else:
                        break
                headers = [
                    [to_html_string(cell_to_obj(row[0]))] for row in table_data[:header_rows]
                ]
                rows = [
                    [to_html_string(cell_to_obj(cell)) for cell in row[1:]] for row in table_data
                ]
                data = rows
            else:
                last_caption = None
                continue
            tableKey = f"table-{table_number}"
            tables[table_index] = {
                "tableNumber": table_number,
                "tableKey": tableKey,
                "caption": last_caption,
                "alt": last_caption,
                "tableType": table_type,
                "headers": headers,
                "hasSpanningCells": has_spanning_cells,
                "data": data,
            }
            table_number += 1
            table_index += 1
            last_caption = None
    return tables, table_number - 1


def handle_table_references(text):
    import re

    def replacer(match):
        table_number = int(match.group(1))
        tableKey = f"table-{str(table_number)}"
        if tableKey:
            return f'<TableReference tableKey="{tableKey}" />'
        else:
            return match.group(0)

    return re.sub(r"Table (\d+)", replacer, text)
